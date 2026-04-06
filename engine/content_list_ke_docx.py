"""
engine/content_list_ke_docx.py — Konverter MinerU content_list.json → DOCX (v4)
================================================================================
Menggunakan metadata struktural MinerU untuk DOCX semirip mungkin dengan PDF:
  - Margin dihitung otomatis dari middle.json (bbox median per halaman)
  - Font: Times New Roman 12pt (standar SP/dokumen pemerintah Indonesia)
  - Indentasi: percentile binning 6 level dari bbox.x0 per dokumen
  - Page break tepat per halaman tanpa spasi ganda
  - Tabel dengan colspan/rowspan penuh
  - Gambar/stempel embedded

Cara pakai:
    from engine.content_list_ke_docx import content_list_ke_bytes
"""

import json
import io
import re
import html as html_lib
import statistics
from pathlib import Path
from typing import Optional, Callable

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────
# KONSTANTA FONT
# ─────────────────────────────────────────────

FONT_BODY = "Times New Roman"
FONT_SIZE_BODY = 12       # pt
FONT_SIZE_HEADING1 = 13   # pt — sedikit lebih besar dari body
FONT_SIZE_HEADING2 = 12   # pt — sama dengan body, tapi bold


# ─────────────────────────────────────────────
# BACA MARGIN DARI middle.json
# ─────────────────────────────────────────────

def _baca_margin_dari_middle(middle_json_path: Path) -> dict:
    """
    Hitung margin halaman dari bbox blok teks di middle.json.
    Return dict: {left, right, top, bottom} dalam Cm.
    Fallback ke margin standar jika middle.json tidak ada.
    """
    FALLBACK = {"left": Cm(3.0), "right": Cm(2.5), "top": Cm(2.5), "bottom": Cm(3.0)}
    if not middle_json_path.exists():
        return FALLBACK

    try:
        with open(middle_json_path, encoding="utf-8") as f:
            mid = json.load(f)

        pages = mid.get("pdf_info", [])
        if not pages:
            return FALLBACK

        W, H = pages[0].get("page_size", [595, 842])
        x0_list, x1_list, y0_list, y1_list = [], [], [], []

        for pg in pages:
            blks = [b for b in pg.get("para_blocks", [])
                    if b.get("type") in ("text", "title") and b.get("bbox")]
            if not blks:
                continue
            bboxes = [b["bbox"] for b in blks]
            x0_list.append(min(b[0] for b in bboxes))
            x1_list.append(max(b[2] for b in bboxes))
            y0_list.append(min(b[1] for b in bboxes))
            y1_list.append(max(b[3] for b in bboxes))

        if not x0_list:
            return FALLBACK

        def pts_ke_cm(pts: float) -> object:
            return Cm(max(1.5, pts / 72 * 2.54))

        ml = statistics.median(x0_list)
        mr = W - statistics.median(x1_list)
        mt = statistics.median(y0_list)
        mb = H - statistics.median(y1_list)

        return {
            "left":   pts_ke_cm(ml),
            "right":  pts_ke_cm(mr),
            "top":    pts_ke_cm(mt),
            "bottom": pts_ke_cm(mb),
        }
    except Exception:
        return FALLBACK


# ─────────────────────────────────────────────
# DETEKSI LEVEL INDENT — Hybrid (Prefix + X0)
# ─────────────────────────────────────────────

# Pola penomoran dokumen legal Indonesia (dari terdalam ke terluar)
_POLA_LEVEL = [
    (re.compile(r"^\(\d{1,2}\)\s"), 4),   # (1) (2) — terdalam
    (re.compile(r"^\([a-z]\)\s"),   3),   # (a) (b)
    (re.compile(r"^\d{1,2}\)\s"),   2),   # 1) 2) (angka + tutup kurung)
    (re.compile(r"^[a-z]\.\s"),     1),   # a. b. (huruf kecil + titik)
]


def _hitung_x0_base(blocks, batas_kanan: float = 350) -> float:
    """
    Hitung x0 baseline (margin kiri dokumen) dari blok body text.
    Ambil percentile 20% dari x0 yang tidak cocok pola penomoran,
    agar tidak terpengaruh outlier OCR atau teks indented.
    """
    FALLBACK = 150.0
    x0_body = []
    for b in blocks:
        if b.get("type") != "text" or not b.get("bbox"):
            continue
        x0 = b["bbox"][0]
        if x0 >= batas_kanan:
            continue
        teks = b.get("text", "").strip()
        # Hanya ambil teks yang TIDAK punya prefix penomoran (body text murni)
        if not any(p.match(teks) for p, _ in _POLA_LEVEL):
            x0_body.append(x0)

    if not x0_body:
        return FALLBACK

    x0_sorted = sorted(x0_body)
    n = len(x0_sorted)
    # Ambil percentile 20% → mendekati margin kiri asli dokumen
    return x0_sorted[max(0, int(n * 0.20))]


def _deteksi_level_hybrid(teks: str, x0: float, x0_base: float) -> int:
    """
    Deteksi level indent: prefix teks (primary) atau x0 offset (fallback).

    1. Jika teks cocok pola penomoran (a., 1), (a), (1)) → gunakan level pola
    2. Jika tidak → hitung offset dari x0_base, petakan ke level 0-4

    Offset thresholds (calibrated dari SP Dhayfullah):
      <20  → level 0 (body)
      20-45 → level 1 (setara a.)
      45-65 → level 2 (setara 1))
      65-90 → level 3 (setara (a))
      ≥90   → level 4 (setara (1))
    """
    stripped = teks.strip()
    for pola, level in _POLA_LEVEL:
        if pola.match(stripped):
            return level

    # Fallback: x0 offset dari baseline margin kiri
    offset = x0 - x0_base
    if offset < 20:
        return 0
    if offset < 45:
        return 1
    if offset < 65:
        return 2
    if offset < 90:
        return 3
    return 4


# ─────────────────────────────────────────────
# SET FONT DEFAULT DOKUMEN
# ─────────────────────────────────────────────

def _set_font_dokumen(doc: Document):
    """Set font default dokumen ke Times New Roman 12pt."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_BODY
    font.size = Pt(FONT_SIZE_BODY)

    # Set juga di elemen XML untuk kompatibilitas penuh
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_BODY)
    rFonts.set(qn("w:hAnsi"), FONT_BODY)
    rFonts.set(qn("w:cs"), FONT_BODY)
    rPr.insert(0, rFonts)

    # Heading 1 — bold, 13pt
    for h_name, h_size in [("Heading 1", FONT_SIZE_HEADING1), ("Heading 2", FONT_SIZE_HEADING2)]:
        if h_name in doc.styles:
            hstyle = doc.styles[h_name]
            hstyle.font.name = FONT_BODY
            hstyle.font.size = Pt(h_size)
            hstyle.font.bold = True
            hstyle.font.color.rgb = None  # hitam (reset dari biru default Word)


def _set_run_font(run):
    """Set font pada sebuah run ke Times New Roman."""
    run.font.name = FONT_BODY
    run.font.size = Pt(FONT_SIZE_BODY)
    # XML-level untuk kompatibilitas
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT_BODY)
    rFonts.set(qn("w:hAnsi"), FONT_BODY)


# ─────────────────────────────────────────────
# FUNGSI UTAMA
# ─────────────────────────────────────────────

def content_list_ke_docx(json_path: Path, judul: str = "") -> Document:
    """
    Konversi *_content_list.json MinerU ke python-docx Document.

    Secara otomatis mencari *_middle.json di folder yang sama untuk
    menghitung margin halaman asli. Font default: Times New Roman 12pt.
    """
    with open(json_path, encoding="utf-8") as f:
        blocks = json.load(f)

    doc = Document()
    _set_font_dokumen(doc)

    # ── Margin dari middle.json ──────────────
    stem = json_path.stem.replace("_content_list", "")
    middle_path = json_path.parent / f"{stem}_middle.json"
    margin = _baca_margin_dari_middle(middle_path)

    for section in doc.sections:
        section.left_margin   = margin["left"]
        section.right_margin  = margin["right"]
        section.top_margin    = margin["top"]
        section.bottom_margin = margin["bottom"]

    # ── Judul opsional ───────────────────────
    if judul:
        p = doc.add_heading(judul, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            _set_run_font(run)

    # ── Hitung x0 baseline untuk deteksi indent ──
    x0_base = _hitung_x0_base(blocks)
    CM_PER_LEVEL = 0.6   # indent Word per level (lebih rapat = lebih mirip PDF)

    halaman_sebelumnya = None
    needs_page_break = False   # flag: halaman baru, pasang page_break_before pada paragraf berikutnya
    pertama_di_halaman = True  # flag: suppress space_before pada elemen pertama halaman baru

    for block in blocks:
        tipe = block.get("type", "")
        page_idx = block.get("page_idx")

        # ── Deteksi pindah halaman ───────────
        if page_idx is not None and halaman_sebelumnya is not None:
            try:
                if int(page_idx) != int(halaman_sebelumnya):
                    needs_page_break = True
                    pertama_di_halaman = True
            except (ValueError, TypeError):
                pass
        if page_idx is not None:
            halaman_sebelumnya = page_idx

        # Skip noise
        if tipe in ("header", "footer", "page_number"):
            continue

        # ─── Blok teks ──────────────────────
        if tipe == "text":
            teks = block.get("text", "").strip()
            if not teks:
                continue

            text_level = block.get("text_level")
            bbox = block.get("bbox", [])
            x0 = bbox[0] if bbox else x0_base
            lvl = _deteksi_level_hybrid(teks, x0, x0_base)

            # Space before: 0 jika blok pertama di halaman baru
            sp_before = Pt(0) if pertama_di_halaman else Pt(2)
            pertama_di_halaman = False

            # Heading level 1 — teks_level=1 ATAU text_level=2 dengan x0 di margin kiri (lvl<=1)
            if text_level == 1:
                p = doc.add_heading(teks, level=1)
                if needs_page_break:
                    p.paragraph_format.page_break_before = True
                    needs_page_break = False
                p.paragraph_format.space_before = sp_before
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.keep_with_next = True   # jangan biarkan heading sendirian di bawah halaman
                for run in p.runs:
                    _set_run_font(run)
                    run.font.bold = True
                continue

            # Heading level 2 — jika text_level=2 DAN tidak punya prefix penomoran
            # (a. b. 1) dll → bukan heading, tapi paragraf indented)
            teks_punya_prefix = any(p.match(teks.strip()) for p, _ in _POLA_LEVEL)
            if text_level == 2 and not teks_punya_prefix:
                p = doc.add_heading(teks, level=2)
                if needs_page_break:
                    p.paragraph_format.page_break_before = True
                    needs_page_break = False
                p.paragraph_format.space_before = sp_before
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.keep_with_next = True   # heading selalu menempel ke konten berikutnya
                for run in p.runs:
                    _set_run_font(run)
                    run.font.bold = True
                continue

            # Paragraf biasa / indented (termasuk text_level=2 dengan lvl>0)
            p = doc.add_paragraph()
            if needs_page_break:
                p.paragraph_format.page_break_before = True
                needs_page_break = False
            if lvl > 0:
                p.paragraph_format.left_indent = Cm(CM_PER_LEVEL * lvl)
            p.paragraph_format.space_before = sp_before
            p.paragraph_format.space_after = Pt(2)
            _tambah_runs(p, teks)

        # ─── Blok tabel ─────────────────────
        elif tipe == "table":
            table_body = block.get("table_body", "")
            caption = block.get("table_caption", "")
            if table_body:
                # Jika perlu page break, sisipkan paragraf kosong sebagai pembawa
                if needs_page_break:
                    pb_para = doc.add_paragraph()
                    pb_para.paragraph_format.page_break_before = True
                    pb_para.paragraph_format.space_before = Pt(0)
                    pb_para.paragraph_format.space_after = Pt(0)
                    needs_page_break = False
                try:
                    _tambah_tabel_html(doc, table_body)
                except Exception:
                    teks_bersih = re.sub(r"<[^>]+>", " ", table_body)
                    teks_bersih = re.sub(r"\s+", " ", teks_bersih).strip()
                    if teks_bersih:
                        doc.add_paragraph(teks_bersih)
            if caption and str(caption) not in ("[]", "", "['']"):
                caption_str = str(caption).strip("[]'\"").strip()
                if caption_str:
                    doc.add_paragraph(caption_str)
            pertama_di_halaman = False

        # ─── Blok gambar / stempel ──────────
        elif tipe in ("image", "seal"):
            img_path_rel = block.get("img_path", "")
            if img_path_rel:
                img_abs = json_path.parent / img_path_rel
                if img_abs.exists():
                    try:
                        p = doc.add_paragraph()
                        if needs_page_break:
                            p.paragraph_format.page_break_before = True
                            needs_page_break = False
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        lebar = Inches(2) if tipe == "seal" else Inches(4)
                        run.add_picture(str(img_abs), width=lebar)
                    except Exception:
                        pass
            pertama_di_halaman = False

    return doc


def content_list_ke_bytes(json_path: Path, judul: str = "") -> bytes:
    """Konversi content_list.json ke bytes DOCX (untuk st.download_button)."""
    doc = content_list_ke_docx(json_path, judul)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def cari_content_list_json(output_dir: Path, stem: str) -> Optional[Path]:
    """Cari *_content_list.json di folder output MinerU."""
    kandidat = [
        output_dir / stem / "auto" / f"{stem}_content_list.json",
        output_dir / stem / "pipeline" / f"{stem}_content_list.json",
    ]
    for p in kandidat:
        if p.exists():
            return p
    hasil = list((output_dir / stem).rglob("*_content_list.json")) if (output_dir / stem).exists() else []
    if not hasil:
        hasil = list(output_dir.rglob("*_content_list.json"))
    return hasil[0] if hasil else None


# ─────────────────────────────────────────────
# HELPER — TABEL dengan colspan/rowspan
# ─────────────────────────────────────────────

def _tambah_tabel_html(doc: Document, html: str):
    """Parse HTML table MinerU (colspan/rowspan) ke Word Table Grid."""
    baris_tr = re.findall(r"<tr[^>]*>(.*?)</tr>", html, re.DOTALL | re.IGNORECASE)
    if not baris_tr:
        return

    parsed_rows = []
    for tr in baris_tr:
        cells_raw = re.findall(
            r"<(td|th)([^>]*)>(.*?)</(?:td|th)>", tr, re.DOTALL | re.IGNORECASE
        )
        row = []
        for tag, attrs, content in cells_raw:
            m_cs = re.search(r'colspan=["\']?(\d+)', attrs, re.I)
            m_rs = re.search(r'rowspan=["\']?(\d+)', attrs, re.I)
            colspan = int(m_cs.group(1)) if m_cs else 1
            rowspan = int(m_rs.group(1)) if m_rs else 1
            teks = re.sub(r"<[^>]+>", " ", content)
            teks = html_lib.unescape(teks)
            teks = re.sub(r"\s+", " ", teks).strip()
            row.append({
                "text": teks,
                "colspan": max(1, colspan),
                "rowspan": max(1, rowspan),
                "header": tag.lower() == "th",
            })
        if row:
            parsed_rows.append(row)

    if not parsed_rows:
        return

    total_cols = max(sum(c["colspan"] for c in row) for row in parsed_rows)
    total_rows = len(parsed_rows)
    if total_cols == 0:
        return

    table = doc.add_table(rows=total_rows, cols=total_cols)
    table.style = "Table Grid"
    occupied = [[False] * total_cols for _ in range(total_rows)]

    for r_idx, row in enumerate(parsed_rows):
        c_grid = 0
        for cell_data in row:
            while c_grid < total_cols and occupied[r_idx][c_grid]:
                c_grid += 1
            if c_grid >= total_cols:
                break

            colspan = cell_data["colspan"]
            rowspan = cell_data["rowspan"]
            word_cell = table.cell(r_idx, c_grid)
            word_cell.text = cell_data["text"][:500]

            # Font dalam sel tabel
            for para in word_cell.paragraphs:
                for run in para.runs:
                    _set_run_font(run)

            if cell_data["header"] or r_idx == 0:
                for para in word_cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

            if colspan > 1 or rowspan > 1:
                end_row = min(r_idx + rowspan - 1, total_rows - 1)
                end_col = min(c_grid + colspan - 1, total_cols - 1)
                if end_row != r_idx or end_col != c_grid:
                    word_cell.merge(table.cell(end_row, end_col))

            for dr in range(rowspan):
                for dc in range(colspan):
                    rr, cc = r_idx + dr, c_grid + dc
                    if rr < total_rows and cc < total_cols:
                        occupied[rr][cc] = True

            c_grid += colspan


# ─────────────────────────────────────────────
# HELPER — INLINE FORMATTING
# ─────────────────────────────────────────────

def _tambah_runs(paragraph, teks: str):
    """Parse **bold**, *italic*, `code` dan tambahkan ke paragraph dengan font TNR."""
    pattern = re.compile(r"(\*\*([^*]+)\*\*|\*([^*]+)\*|`([^`]+)`|([^*`]+))")
    for m in pattern.finditer(teks):
        if m.group(2):
            run = paragraph.add_run(m.group(2))
            _set_run_font(run)
            run.bold = True
        elif m.group(3):
            run = paragraph.add_run(m.group(3))
            _set_run_font(run)
            run.italic = True
        elif m.group(4):
            run = paragraph.add_run(m.group(4))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif m.group(5):
            run = paragraph.add_run(m.group(5))
            _set_run_font(run)
