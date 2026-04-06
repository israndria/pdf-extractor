"""
engine/md_ke_docx.py — Konverter Markdown → DOCX
==================================================
Menggunakan python-docx untuk konversi teks Markdown hasil MinerU
ke format Word (.docx) yang bisa diedit.

Cara pakai:
    from engine.md_ke_docx import konversi_md_ke_docx, md_ke_bytes
"""

import re
import io
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ─────────────────────────────────────────────
# FUNGSI UTAMA
# ─────────────────────────────────────────────

def konversi_md_ke_docx(teks_markdown: str, judul: str = "") -> Document:
    """
    Konversi teks Markdown ke objek python-docx Document.

    Args:
        teks_markdown: Isi teks Markdown dari MinerU
        judul: Judul dokumen (opsional, ditambahkan sebagai heading utama)

    Returns:
        docx.Document object
    """
    doc = Document()

    # Setting margin halaman (2cm semua sisi)
    from docx.shared import Cm
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Judul dokumen (jika ada)
    if judul:
        p = doc.add_heading(judul, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    baris_list = teks_markdown.split("\n")
    i = 0
    while i < len(baris_list):
        baris = baris_list[i]

        # ─── Pemisah halaman (MinerU pakai ---)
        if baris.strip() == "---":
            doc.add_page_break()
            i += 1
            continue

        # ─── Heading (#, ##, ###)
        heading_match = re.match(r"^(#{1,6})\s+(.*)", baris)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            teks_heading = _strip_md(heading_match.group(2))
            doc.add_heading(teks_heading, level=level)
            i += 1
            continue

        # ─── Tabel HTML (dari MinerU — tabel dirender sebagai <table>...)
        if baris.strip().startswith("<table") or baris.strip().startswith("|"):
            # Kumpulkan semua baris tabel
            blok_tabel = []
            while i < len(baris_list):
                baris_tabel = baris_list[i]
                if baris_tabel.strip().startswith("<table") or baris_tabel.strip().startswith("|") or \
                   baris_tabel.strip().startswith("<tr") or baris_tabel.strip().startswith("<td") or \
                   baris_tabel.strip().startswith("<th") or baris_tabel.strip().startswith("</") or \
                   (blok_tabel and baris_tabel.strip() == ""):
                    if baris_tabel.strip() == "" and blok_tabel:
                        break
                    blok_tabel.append(baris_tabel)
                    i += 1
                else:
                    break
            _tambah_tabel(doc, "\n".join(blok_tabel))
            continue

        # ─── List item (-, *, 1.)
        list_match = re.match(r"^(\s*)[-*]\s+(.*)", baris)
        list_num_match = re.match(r"^(\s*)\d+[.)]\s+(.*)", baris)
        if list_match or list_num_match:
            m = list_match or list_num_match
            indent = len(m.group(1)) // 2
            teks_item = _strip_md(m.group(2))
            p = doc.add_paragraph(style="List Bullet" if list_match else "List Number")
            _tambah_runs(p, teks_item)
            i += 1
            continue

        # ─── Baris kosong → spasi
        if baris.strip() == "":
            i += 1
            continue

        # ─── Paragraf biasa
        teks_bersih = _strip_md(baris)
        if teks_bersih:
            p = doc.add_paragraph()
            _tambah_runs(p, teks_bersih)

        i += 1

    return doc


def md_ke_bytes(teks_markdown: str, judul: str = "") -> bytes:
    """
    Konversi Markdown ke bytes DOCX (untuk st.download_button).
    """
    doc = konversi_md_ke_docx(teks_markdown, judul)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ─────────────────────────────────────────────
# HELPER — PARSING INLINE
# ─────────────────────────────────────────────

def _strip_md(teks: str) -> str:
    """Bersihkan karakter markdown sederhana."""
    # Hapus image markdown ![alt](url)
    teks = re.sub(r"!\[[^\]]*\]\([^\)]*\)", "", teks)
    # Hapus link [text](url) → text
    teks = re.sub(r"\[([^\]]+)\]\([^\)]*\)", r"\1", teks)
    return teks.strip()


def _tambah_runs(paragraph, teks: str):
    """
    Parse teks dengan formatting inline (**bold**, *italic*, `code`)
    dan tambahkan ke paragraph sebagai runs.
    """
    # Pattern: **bold**, *italic*, `code`
    pattern = re.compile(r"(\*\*([^*]+)\*\*|\*([^*]+)\*|`([^`]+)`|([^*`]+))")
    for m in pattern.finditer(teks):
        if m.group(2):  # **bold**
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(3):  # *italic*
            run = paragraph.add_run(m.group(3))
            run.italic = True
        elif m.group(4):  # `code`
            run = paragraph.add_run(m.group(4))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif m.group(5):  # plain text
            paragraph.add_run(m.group(5))


def _tambah_tabel(doc: Document, blok: str):
    """
    Parse tabel dari Markdown pipe format atau HTML table, tambahkan ke Document.
    """
    # ─── Format pipe Markdown (| col1 | col2 |)
    baris_pipe = [b for b in blok.split("\n") if b.strip().startswith("|") and b.strip().endswith("|")]
    if baris_pipe:
        # Filter baris separator (|---|---|)
        baris_data = [b for b in baris_pipe if not re.match(r"^\|[\s\-:|]+\|$", b.strip())]
        if not baris_data:
            return
        rows = []
        for b in baris_data:
            cells = [c.strip() for c in b.strip().strip("|").split("|")]
            rows.append(cells)

        max_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                if c_idx < max_cols:
                    cell = table.cell(r_idx, c_idx)
                    cell.text = _strip_md(cell_text)
                    if r_idx == 0:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
        doc.add_paragraph()
        return

    # ─── Format HTML table (dari MinerU)
    if "<table" in blok.lower():
        try:
            _parse_html_table(doc, blok)
        except Exception:
            # Fallback: tulis sebagai teks biasa
            teks_bersih = re.sub(r"<[^>]+>", " ", blok)
            teks_bersih = re.sub(r"\s+", " ", teks_bersih).strip()
            if teks_bersih:
                doc.add_paragraph(teks_bersih)
        doc.add_paragraph()


def _parse_html_table(doc: Document, html: str):
    """Parse HTML table dari MinerU dan tambahkan ke Document."""
    import html as html_lib

    # Ekstrak semua <tr>
    baris_tr = re.findall(r"<tr[^>]*>(.*?)</tr>", html, re.DOTALL | re.IGNORECASE)
    if not baris_tr:
        return

    rows_data = []
    for tr in baris_tr:
        # Ekstrak <td> atau <th>
        cells = re.findall(r"<(?:td|th)[^>]*>(.*?)</(?:td|th)>", tr, re.DOTALL | re.IGNORECASE)
        row = []
        for cell in cells:
            # Hapus tag HTML dalam sel
            teks = re.sub(r"<[^>]+>", " ", cell)
            teks = html_lib.unescape(teks)
            teks = re.sub(r"\s+", " ", teks).strip()
            row.append(teks)
        if row:
            rows_data.append(row)

    if not rows_data:
        return

    max_cols = max(len(r) for r in rows_data)
    if max_cols == 0:
        return

    table = doc.add_table(rows=len(rows_data), cols=max_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows_data):
        for c_idx, cell_text in enumerate(row[:max_cols]):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text[:500]  # batasi 500 karakter per sel
            if r_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
